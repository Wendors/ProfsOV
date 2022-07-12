from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt,Inches
import tempfile
#Створити вордовський файл для завантаженя контролю Який збережиний шаблон docx в ресурсі программи
class Docxs():
    gnam = "/gerb.png"
    pathtemp = tempfile.gettempdir() + "/Proftemp"
    def __init__(self):
        self.document = Document()
        self.p2 = self.document.add_paragraph()
        self.rung = self.p2.add_run()
        self.rung.add_picture('{0}'.format(self.pathtemp + self.gnam), width=Cm(1.3), height=Cm(1.7))
        self.rung.add_break(break_type=WD_BREAK.LINE)
        self.run1 = self.p2.add_run()
        self.run1.add_text('МІНІСТЕРСТВО ЮСТИЦІЇ УКРАЇНИ')
        self.run1.font.name = 'Times New Roman'
        self.run1.font.size = Pt(16)
        self.run1.bold = True
        self.run1.add_break(break_type=WD_BREAK.LINE)
        self.run2 = self.p2.add_run()
        self.run2.add_text('ДЕРЖАВНА УСТАНОВА «ПОЛИЦЬКА ВИПРАВНА КОЛОНІЯ (№76)»')
        self.run2.font.name = 'Times New Roman'
        self.run2.font.size = Pt(14)
        self.run2.bold = True
        self.run2.add_break(break_type=WD_BREAK.LINE)
        self.run3 = self.p2.add_run()
        self.run3.add_text(' с. Іванчі Вараського району Рівненської  області, 34375тел. 2-53-24; 5-31-49,'
                           'тел./факс 5-35-46')
        self.run3.font.name = 'Times New Roman'
        self.run3.add_break(break_type=WD_BREAK.LINE)
        self.run3.add_text('E-mail: pvk76@ukr.net Код ЄДРПОЧ 08564370')
        self.run3.font.size = Pt(11)
        self.run3.add_break(break_type=WD_BREAK.LINE)
        self.p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self.p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        self.p3 = self.document.add_paragraph()
        self.run4 = self.p3.add_run()
        self.run4.add_text('__________№___________')
        self.run4.add_tab()
        self.run4.add_tab()
        self.run4.add_tab()
        self.run4.add_tab()
        self.run4.add_text("На №5к/вих.//7709 від 25.10.2021")
        self.run4.font.name = 'Times New Roman'
        self.run4.font.size = Pt(14)
        self.p3.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        self.pa1 = self.document.add_paragraph()
        self.pa1.paragraph_format.left_indent = Cm(9)
        self.rpa1 = self.pa1.add_run()
        self.rpa1.font.size = Pt(14)
        self.rpa1.bold = True
        self.rpa1.add_text("Начальнику Західного міжрегіонального управління з"
                           " питань виконання кримінальних покарань Міністерства юстиції")
        self.rpa1.add_break(break_type=WD_BREAK.LINE)
        self.rpa1.add_text("полковнику внутрішньої служби")
        self.rpa1.add_break(break_type=WD_BREAK.LINE)
        self.rpa1.add_text("Ярославу БОЙКУ")
        self.pa1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        self.pa1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        self.pa1.paragraph_format.space_after = Cm(0.5)
        self.pa2 = self.document.add_paragraph()
        self.rpa2 = self.pa2.add_run()
        self.rpa2.font.size = Pt(14)
        self.rpa2.add_text("Щодо оперативно-профілактичного обліку")
        self.pa2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        self.pa2.paragraph_format.space_after = Cm(0.5)
        self.psh = self.document.add_paragraph()
        self.rsh = self.psh.add_run()
        self.rsh.add_text('Шановний  Ярославе Володимировичу!')
        self.rsh.font.name = 'Times New Roman'
        self.rsh.font.size = Pt(14)
        self.rsh.bold = True
        self.psh.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self.ptext = self.document.add_paragraph()
        self.rtext = self.ptext.add_run()
        self.rtext.font.name = 'Times New Roman'
        self.rtext.font.size = Pt(14)
        self.rtext.add_tab()
        self.rtext.add_text("На виконання вказівки Західного міжрегіонального управління"
                            " з питань виконання кримінальних покарань Міністерства юстиції"
                            " №5к/вих.//7709 від 25.10.2021р., та з метою забезпечення безпосереднього"
                            " контролю за поведінкою осіб, які мають стійку антисоціальну орієнтацію,"
                            " вчинили тяжкі резонансні злочини або схильні до їх вчинення, а також"
                            " своєчасного проведення стосовно них відповідних заходів щодо профілактики"
                            " кримінальних правопорушень, надсилаю Вам списки зазначеної категорії осіб"
                            " за червень місяць 2022 року")
        self.ptext.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        self.ptext.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        self.ptext.paragraph_format.space_after = Cm(0.5)
        self.pd = self.document.add_paragraph()
        self.raut = self.pd.add_run()
        self.raut.font.name = 'Times New Roman'
        self.raut.font.size = Pt(14)
        self.raut.add_text("Додаток на 1 арк.")
        self.pd.paragraph_format.space_after = Cm(0.5)
        self.pdu = self.document.add_paragraph()
        self.rdo = self.pdu.add_run()
        self.rdo.font.name = 'Times New Roman'
        self.rdo.font.size = Pt(14)
        self.rdo.add_text("З повагою,")
        self.paut = self.document.add_paragraph()
        self.raut = self.paut.add_run()
        self.raut.font.name = 'Times New Roman'
        self.raut.font.size = Pt(14)
        self.raut.bold = True
        self.raut.add_text("В. о. начальника державної установи")
        self.raut.add_break(break_type=WD_BREAK.LINE)
        self.raut.add_text("«Полицька виправна колонія (№76)»")
        self.raut.add_break(break_type=WD_BREAK.LINE)
        self.raut.add_text("підполковник внутрішньої служби")
        self.raut.add_tab()
        self.raut.add_tab()
        self.raut.add_tab()
        self.raut.add_text("Руслан БОГДАШЕНКО")
        self.raut.add_break(break_type=WD_BREAK.LINE)
        self.raut.add_break(break_type=WD_BREAK.LINE)
        self.paut.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        self.paut.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        self.rauts = self.paut.add_run()
        self.rauts.font.name = 'Times New Roman'
        self.rauts.font.size = Pt(10)
        self.rauts.add_text("Сергій ПОЛУНЕЦЬ")
        self.rauts.add_break(break_type=WD_BREAK.LINE)
        self.rauts.add_text("тел. 09837334329")

        self.document.save('demo.docx')
Docxs()