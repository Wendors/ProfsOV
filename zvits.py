# -*- coding: utf-8 -*-

import os
import sys
import tempfile
import pickle
from Datas import Dat
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt,Inches
from PyQt5 import QtCore, QtGui, QtWidgets

if 'ANDROID_BOOTLOGO' in os.environ:
    pass
else:
    from PyQt5 import QtPrintSupport
import Profs_rc

class Window(QtWidgets.QWidget):
    pathtemp = tempfile.gettempdir() + "/Proftemp"
    orient = 0
    wor = 0
    cur = 0
    def __init__(self, titles):
        QtWidgets.QDialog.__init__(self)
        self.setFixedSize(960, 600)
        self.resolution = QtWidgets.QDesktopWidget().screenGeometry()
        if 'ANDROID_BOOTLOGO' in os.environ:
            self.move(int((self.resolution.width() / 2) - (self.frameSize().width() / 2)), int(0))
        else:
            self.move(int((self.resolution.width() / 2) - (self.frameSize().width() / 2)),
                          int((self.resolution.height() / 2) - (self.frameSize().height() / 2)))
        self.activateWindow()
        self.setWindowTitle(self.tr(titles))
        icon = QtGui.QIcon()
        icon.addPixmap(QtGui.QPixmap(":/Icon/Profico.ico"), QtGui.QIcon.Normal, QtGui.QIcon.Off)
        self.setWindowModality(QtCore.Qt.ApplicationModal)
        self.setWindowFlags(QtCore.Qt.WindowCloseButtonHint)
        self.setWindowIcon(icon)
        self.editor = QtWidgets.QTextEdit(self)
        self.editor.setLineWrapMode(self.editor.WidgetWidth)
        self.editor.textChanged.connect(self.handleTextChanged)
        self.buttonOpen = QtWidgets.QPushButton('В pdf', self)
        self.buttonOpen.setAutoDefault(True)
        self.buttonOpen.clicked.connect(self.topdf)
        self.buttonPrint = QtWidgets.QPushButton('Друкувати', self)
        self.buttonPrint.setAutoDefault(True)
        self.buttonPrint.clicked.connect(self.handlePrint)
        self.buttonPreview = QtWidgets.QPushButton('Переглянути', self)
        self.buttonPreview.setAutoDefault(True)
        self.buttonPreview.clicked.connect(self.handlePreview)
        self.buttonCurs = QtWidgets.QPushButton('В word', self)
        self.buttonCurs.setAutoDefault(True)
        self.buttonCurs.clicked.connect(self.hirCurs)
        layout = QtWidgets.QGridLayout(self)
        layout.addWidget(self.editor, 0, 0, 1, 4)
        layout.addWidget(self.buttonCurs,1, 0)
        layout.addWidget(self.buttonOpen, 1, 1)
        layout.addWidget(self.buttonPrint, 1, 2)
        layout.addWidget(self.buttonPreview, 1, 3)
        self.handleTextChanged()
        self.setTabOrder(self.buttonCurs,self.buttonPreview)
        self.setTabOrder(self.buttonPreview, self.buttonPrint)
        self.setTabOrder(self.buttonPrint, self.buttonOpen)
        self.setTabOrder(self.buttonOpen, self.editor)

    def hirCurs(self):
        self.Wordtext()


    def handleOpen(self, orent, wor):
        self.orient = orent
        self.wor = wor
        path = self.pathtemp + "/_temp.html"
        info = QtCore.QFileInfo(path)
        f = open(path, "r")
        text = f.read()
        f.close()
        if self.wor == 1:
            self.buttonCurs.setEnabled(False)
        if info.completeSuffix() == 'html':
            self.editor.setHtml(text)
        else:
            self.editor.setPlainText(text)

    def handlePrint(self):
        location = self.pathtemp + "/_temp.html"
        with open(location,"w") as f:
            f.write(self.editor.document().toHtml())
            f.close()
        icon = QtGui.QIcon()
        icon.addPixmap(QtGui.QPixmap(":/Icon/Profico.ico"), QtGui.QIcon.Normal, QtGui.QIcon.Off)
        doc = QtGui.QTextDocument()
        location = self.pathtemp + "/_temp.html"
        html = open(location).read()
        doc.setHtml(html)
        if 'ANDROID_BOOTLOGO' in os.environ:
            pass
        else:
            printerd = QtPrintSupport.QPrinter(QtPrintSupport.QPrinter.HighResolution)
            if self.orient == 1:
                printerd.setOrientation(QtPrintSupport.QPrinter.Landscape)
                printerd.setResolution(300)
                printerd.setPageMargins(30,10,10,20, QtPrintSupport.QPrinter.Millimeter)
                printerd.setPageSize(QtPrintSupport.QPrinter.A4)
                sizes = QtCore.QSizeF(794, 1123)
            else:
                printerd.setOrientation(QtPrintSupport.QPrinter.Portrait)
                printerd.setResolution(300)
                printerd.setPageMargins(30,10,10,20, QtPrintSupport.QPrinter.Millimeter)
                printerd.setPageSize(QtPrintSupport.QPrinter.A4)
                sizes = QtCore.QSizeF(794, 1123)
            dialog = QtPrintSupport.QPrintDialog(printerd,self)
            dialog.setWindowFlags(QtCore.Qt.Dialog | QtCore.Qt.WindowSystemMenuHint | QtCore.Qt.WindowCloseButtonHint)
            dialog.setWindowIcon(icon)
            doc.setPageSize(sizes)
            if dialog.exec_() == QtWidgets.QDialog.Accepted:
                doc.print_(dialog.printer())

    def handlePreview(self):
        location = self.pathtemp + "/_temp.html"
        with open(location,"w") as f:
            f.write(self.editor.document().toHtml())
            f.close()
        icon = QtGui.QIcon()
        icon.addPixmap(QtGui.QPixmap(":/Icon/Profico.ico"), QtGui.QIcon.Normal, QtGui.QIcon.Off)
        doc = QtGui.QTextDocument()
        location = self.pathtemp + "/_temp.html"
        html = open(location).read()
        doc.setHtml(html)
        if 'ANDROID_BOOTLOGO' in os.environ:
            pass
        else:
            printerd = QtPrintSupport.QPrinter(QtPrintSupport.QPrinter.PrinterResolution)
            if self.orient == 1:
                printerd.setOrientation(QtPrintSupport.QPrinter.Landscape)
                printerd.setResolution(300)
                printerd.setPageSize(QtPrintSupport.QPrinter.A4)
                printerd.setPageMargins(30,10,10,20, QtPrintSupport.QPrinter.Millimeter)
                sizes = QtCore.QSizeF(1123, 794)
            else:
                printerd.setOrientation(QtPrintSupport.QPrinter.Portrait)
                printerd.setResolution(300)
                printerd.setPageSize(QtPrintSupport.QPrinter.A4)
                printerd.setPageMargins(30,10,10,20, QtPrintSupport.QPrinter.Millimeter)
                sizes = QtCore.QSizeF(794, 1123)
            dialog = QtPrintSupport.QPrintPreviewDialog(printerd,self)
            dialog.setWindowFlags(QtCore.Qt.Dialog | QtCore.Qt.WindowSystemMenuHint | QtCore.Qt.WindowCloseButtonHint)
            dialog.setWindowIcon(icon)
            doc.setPageSize(sizes)
            dialog.paintRequested.connect(doc.print_)
            dialog.exec_()

    def handleTextChanged(self):
        enable = not self.editor.document().isEmpty()
        self.buttonPrint.setEnabled(enable)
        self.buttonPreview.setEnabled(enable)
        self.buttonOpen.setEnabled(enable)

    def topdf(self):
        if sys.platform == 'win32':
            self.sevs = os.environ['USERPROFILE'] + "/Documents"
        if sys.platform == 'linux':
            self.sevs = os.environ['HOME']
        self.files = QtWidgets.QFileDialog.getSaveFileName(self, "Зберегти дані", self.sevs, "PDF (*.pdf)")
        self.files = self.files[0]
        if self.files != "":
            doc = QtGui.QTextDocument()
            location = self.pathtemp + "/_temp.html"
            with open(location, "w") as f:
                f.write(self.editor.document().toHtml())
                f.close()
            html = open(location).read()
            doc.setHtml(html)
            if 'ANDROID_BOOTLOGO' in os.environ:
                pass
            else:
                printer = QtPrintSupport.QPrinter()
                printer.setOutputFileName(self.files)
                printer.setOutputFormat(QtPrintSupport.QPrinter.PdfFormat)
                if self.orient == 1:
                    printer.setOrientation(QtPrintSupport.QPrinter.Landscape)
                    printer.setResolution(300)
                    printer.setPageSize(QtPrintSupport.QPrinter.A4)
                    printer.setPageMargins(30, 10, 10, 20, QtPrintSupport.QPrinter.Millimeter)
                    sizes = QtCore.QSizeF(1123, 794)
                else:
                    printer.setOrientation(QtPrintSupport.QPrinter.Portrait)
                    printer.setResolution(300)
                    printer.setPageSize(QtPrintSupport.QPrinter.A4)
                    printer.setPageMargins(30, 10, 10, 20, QtPrintSupport.QPrinter.Millimeter)
                    sizes = QtCore.QSizeF(794, 1123)
                doc.setPageSize(sizes)
                doc.print_(printer)
    def Wordtext(self):
        self.gnam = "/gerb.png"
        self.word = "/.default.docx"
        self.images = QtGui.QImage(":gerb/Gerb.png")
        self.images.save(self.pathtemp + self.gnam)
        QtCore.QFile.copy(":/Docx/default.docx",self.pathtemp + self.word)
        self.mun, self.years = Dat().Daetes()
        self.document = Document(docx=self.pathtemp + self.word)
        self.p2 = self.document.add_paragraph()
        self.rung = self.p2.add_run()
        self.rung.add_picture('{0}'.format(self.pathtemp + self.gnam), width=Cm(1.3), height=Cm(1.7))
        self.rung.add_break(break_type=WD_BREAK.LINE)
        self.run1 = self.p2.add_run()
        self.run1.add_text('МІНІСТЕРСТВО ЮСТИЦІЇ УКРАЇНИ')
        self.run1.font.name = 'Times New Roman'
        self.run1.font.size = Pt(16)
        self.run1.bold = True
        self.run1.add_break(break_type=WD_BREAK.LINE)
        self.run2 = self.p2.add_run()
        self.run2.add_text('ДЕРЖАВНА УСТАНОВА «ПОЛИЦЬКА ВИПРАВНА КОЛОНІЯ (№76)»')
        self.run2.font.name = 'Times New Roman'
        self.run2.font.size = Pt(14)
        self.run2.bold = True
        self.run2.add_break(break_type=WD_BREAK.LINE)
        self.run3 = self.p2.add_run()
        self.run3.add_text(' с. Іванчі Вараського району Рівненської  області, 34375тел. 2-53-24; 5-31-49,'
                           'тел./факс 5-35-46')
        self.run3.font.name = 'Times New Roman'
        self.run3.add_break(break_type=WD_BREAK.LINE)
        self.run3.add_text('E-mail: pvk76@ukr.net Код ЄДРПОУ 08564370')
        self.run3.font.size = Pt(11)
        self.run3.add_break(break_type=WD_BREAK.LINE)
        self.p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self.p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        self.p3 = self.document.add_paragraph()
        self.run4 = self.p3.add_run()
        self.run4.add_text('__________№___________')
        self.run4.add_tab()
        self.run4.add_tab()
        self.run4.add_tab()
        self.run4.add_text("На №5к/вих.//7709 від 25.10.2021")
        self.run4.font.name = 'Times New Roman'
        self.run4.font.size = Pt(14)
        self.p3.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        self.pa1 = self.document.add_paragraph()
        self.pa1.paragraph_format.left_indent = Cm(9)
        self.rpa1 = self.pa1.add_run()
        self.rpa1.font.name = 'Times New Roman'
        self.rpa1.font.size = Pt(14)
        self.rpa1.bold = True
        self.rpa1.add_text("Начальнику Західного міжрегіонального управління з"
                           " питань виконання кримінальних покарань Міністерства юстиції")
        self.rpa1.add_break(break_type=WD_BREAK.LINE)
        self.rpa1.add_text("капітану внутрішньої служби")
        self.rpa1.add_break(break_type=WD_BREAK.LINE)
        self.rpa1.add_text("Ігорю ТИМОЧКУ")
        self.pa1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        self.pa1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        self.pa1.paragraph_format.space_after = Cm(0.5)
        self.pa2 = self.document.add_paragraph()
        self.rpa2 = self.pa2.add_run()
        self.rpa2.font.name = 'Times New Roman'
        self.rpa2.font.size = Pt(14)
        self.rpa2.add_text("Щодо оперативно-профілактичного обліку")
        self.pa2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        self.pa2.paragraph_format.space_after = Cm(0.5)
        self.psh = self.document.add_paragraph()
        self.rsh = self.psh.add_run()
        self.rsh.add_text('Шановний Ігорю Ярославовичу!')
        self.rsh.font.name = 'Times New Roman'
        self.rsh.font.size = Pt(14)
        self.rsh.bold = True
        self.psh.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self.ptext = self.document.add_paragraph()
        self.rtext = self.ptext.add_run()
        self.rtext.font.name = 'Times New Roman'
        self.rtext.font.size = Pt(14)
        self.rtext.add_tab()
        self.rtext.add_text("На виконання вказівки Західного міжрегіонального управління"
                            " з питань виконання кримінальних покарань Міністерства юстиції"
                            " №5к/вих.//7709 від 25.10.2021р., та з метою забезпечення безпосереднього"
                            " контролю за поведінкою осіб, які мають стійку антисоціальну орієнтацію,"
                            " вчинили тяжкі резонансні злочини або схильні до їх вчинення, а також"
                            " своєчасного проведення стосовно них відповідних заходів щодо профілактики"
                            " кримінальних правопорушень, надсилаю Вам списки зазначеної категорії осіб"
                            " за {0} місяць {1} року".format(self.mun,self.years))
        self.ptext.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        self.ptext.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        self.ptext.paragraph_format.space_after = Cm(0.5)
        self.pd = self.document.add_paragraph()
        self.raut = self.pd.add_run()
        self.raut.font.name = 'Times New Roman'
        self.raut.font.size = Pt(14)
        self.raut.add_text("Додаток на 1 арк.")
        self.pd.paragraph_format.space_after = Cm(0.5)
        self.pdu = self.document.add_paragraph()
        self.rdo = self.pdu.add_run()
        self.rdo.font.name = 'Times New Roman'
        self.rdo.font.size = Pt(14)
        self.rdo.add_text("З повагою,")
        self.paut = self.document.add_paragraph()
        self.raut = self.paut.add_run()
        self.raut.font.name = 'Times New Roman'
        self.raut.font.size = Pt(14)
        self.raut.bold = True
        self.raut.add_text("Начальник державної установи")
        self.raut.add_break(break_type=WD_BREAK.LINE)
        self.raut.add_text("«Полицька виправна колонія (№76)»")
        self.raut.add_break(break_type=WD_BREAK.LINE)
        self.raut.add_text("полковник внутрішньої служби")
        self.raut.add_tab()
        self.raut.add_tab()
        self.raut.add_tab()
        self.raut.add_tab()
        self.raut.add_text("Сергій ГРАБОВСЬКИЙ")
        self.raut.add_break(break_type=WD_BREAK.LINE)
        self.raut.add_break(break_type=WD_BREAK.LINE)
        self.paut.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        self.paut.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        self.rauts = self.paut.add_run()
        self.rauts.font.name = 'Times New Roman'
        self.rauts.font.size = Pt(10)
        self.rauts.add_text("Сергій ПОЛУНЕЦЬ")
        self.rauts.add_break(break_type=WD_BREAK.LINE)
        self.rauts.add_text("тел. 09837334329")
        self.document.add_page_break()
        self.Worrd_2()
    def Worrd_2(self):
        self.pol = self.document.add_paragraph()
        self.pol.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        self.rpol = self.pol.add_run()
        self.rpol.add_text('Державна установа')
        self.rpol.font.name = 'Times New Roman'
        self.rpol.font.size = Pt(16)
        self.rpol.bold = True
        self.rpol.add_break(break_type=WD_BREAK.LINE)
        self.rpol_1 = self.pol.add_run()
        self.rpol_1.add_text('«Полицька виправна колонія (№76)»')
        self.rpol_1.font.name = 'Times New Roman'
        self.rpol_1.font.size = Pt(16)
        self.rpol_1.bold = True
        self.pol.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self.ffs = open(self.pathtemp + "/Profs.dbsp", "r")
        self.filess = self.ffs.read()
        self.ffs.close()
        try:
            self.listprof = Dat().Spis_profs()
            self.opfils = open(self.filess, "rb")
            self.datas = pickle.load(self.opfils)
            self.opfils.close()
            self.proerss = self.datas
            self.plist = []
            for i in self.listprof:
                if i == "Напад":
                    pass
                else:
                    for ip in self.proerss:
                        self.i = self.proerss.get(i)
                        if str(ip) == str(i):
                            self.p_ob = self.document.add_paragraph()
                            self.p_ob.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                            self.r_ob = self.p_ob.add_run()
                            self.r_ob.add_text('{0}'.format(str(self.listprof.get(ip))))
                            self.r_ob.font.name = 'Times New Roman'
                            self.r_ob.font.size = Pt(14)
                            self.r_ob.bold = True
                            self.p_ob.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            self.p_ob.paragraph_format.space_after = Pt(1)
                            self.p_ob.paragraph_format.space_before = Pt(1)
                            self.p_obp = self.document.add_paragraph()
                            self.p_obp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                            self.r_obp = self.p_obp.add_run()
                            if self.i.__len__() > 0:
                                self.list_keys = list(self.i.keys())
                                self.list_keys.sort()
                                self.xx = 0
                                for ia in self.list_keys:
                                    self.xx += 1
                                    self.priis = self.i.setdefault(ia)
                                    self._soname = str("{0}".format(self.priis['soname']))
                                    self._name = str("{0}".format(self.priis['name']))
                                    self._father = str("{0}".format(self.priis['father']))
                                    self._brsd = str("{0}".format(self.priis['birsdey']))
                                    self.r_obp.font.name = 'Times New Roman'
                                    self.r_obp.font.size = Pt(14)
                                    self.r_obp.add_text('{0}. {1} {2} {3} {4} р.н.'.format(self.xx, self._soname, self._name, self._father, self._brsd))
                                    if self.list_keys.__len__() != self.xx:
                                        self.r_obp.add_break(break_type=WD_BREAK.LINE)
                                    self.p_obp.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            self.p_ob_p = self.document.add_paragraph()
            self.p_ob_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            self.r_ob_p = self.p_ob_p.add_run()
            self.r_ob_p.add_text('Оперуповноважений оперативного відділу')
            self.r_ob_p.add_break(break_type=WD_BREAK.LINE)
            self.r_ob_p.add_text('державної установи «Полицька')
            self.r_ob_p.add_break(break_type=WD_BREAK.LINE)
            self.r_ob_p.add_text('виправна колонія (№76)»')
            self.r_ob_p.font.name = 'Times New Roman'
            self.r_ob_p.add_break(break_type=WD_BREAK.LINE)
            self.r_ob_p.add_text('капітан внутрішньої служби')
            self.r_ob_p.add_tab()
            self.r_ob_p.add_tab()
            self.r_ob_p.add_tab()
            self.r_ob_p.add_tab()
            self.r_ob_p.add_text('Сергій ПОЛУНЕЦЬ')
            self.r_ob_p.font.size = Pt(14)
            self.r_ob_p.bold = True
        except:
            pass
        if sys.platform == 'win32':
            self.document.save(os.environ['USERPROFILE'] + self.word)
            os.startfile(os.environ['USERPROFILE'] + self.word)
        if sys.platform == 'linux':
            self.document.save(os.environ['HOME'] + self.word)
            os.system("xdg-open {0}".format(os.environ['HOME'] + self.word))
        self.close()
